"""文件导出模块 — CSV / XLSX / Markdown / HTML / PDF / DOCX 文件生成。

支持任务隔离：每个 task_id 在 outputs/{task_id}/ 独立子目录下生成文件。"""

import csv
import logging
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)

_BASE_OUTPUT_DIR = Path(__file__).parent.parent / "outputs"

HEADERS = ["序号", "姓名", "邮箱", "学院", "职称", "主页链接"]


def get_task_dir(task_id: str = "") -> Path:
    """返回任务专属输出目录。task_id 为空时使用兼容根目录。"""
    if task_id:
        safe = task_id.replace("/", "_").replace("\\", "_").replace("..", "_")
        d = _BASE_OUTPUT_DIR / safe
    else:
        d = _BASE_OUTPUT_DIR
    d.mkdir(parents=True, exist_ok=True)
    return d


def cleanup_task_dir(task_id: str) -> None:
    """删除任务专属目录及其所有文件。"""
    if not task_id:
        return
    d = _BASE_OUTPUT_DIR / task_id.replace("/", "_").replace("\\", "_").replace("..", "_")
    if not d.exists():
        return
    for f in d.iterdir():
        try:
            f.unlink()
        except Exception:
            pass
    try:
        d.rmdir()
    except Exception:
        pass


def _make_filename(university: str, ext: str, task_id: str = "") -> str:
    import re
    # 1. 尝试从任务历史提取自定义需求 (如：院系/专业分类等名称)
    requirements = ""
    if task_id:
        try:
            from agent.history import history
            task = history.get(task_id)
            if task:
                user_msg = ""
                for m in task.get("messages", []):
                    if m.get("role") == "user":
                        user_msg = m.get("content", "")
                        break
                if not user_msg:
                    user_msg = task.get("title", "")
                
                # 匹配院/系/中心/实验室/研究院等标志性层级名词
                m = re.search(r"([一-鿿]{2,15}(?:学院|系|中心|研究所|研究室|实验室|部))", user_msg)
                if m:
                    dept = m.group(1)
                    if university in dept:
                        dept = dept.replace(university, "")
                    requirements = f"_{dept}"
        except Exception as e:
            logger.warning(f"提取自定义需求失败: {e}")

    safe_uni = (university or "unknown").replace(" ", "_")
    
    # 2. 动态检测已有文件版本，实现 V X.X.X 的自动增长 (从 1.0.0 开始递增最后一位)
    prefix = f"{safe_uni}{requirements}_V"
    max_major, max_minor, max_patch = 1, 0, 0
    found = False
    try:
        base = _BASE_OUTPUT_DIR.resolve()
        for f in base.rglob(f"*{ext}"):
            name = f.name
            if name.startswith(prefix):
                # 截取版本号部分 (e.g. major.minor.patch)
                ver_part = name[len(prefix):-len(ext)-1]
                m = re.match(r"^(\d+)\.(\d+)\.(\d+)", ver_part)
                if m:
                    found = True
                    major, minor, patch = map(int, m.groups())
                    if (major, minor, patch) > (max_major, max_minor, max_patch):
                        max_major, max_minor, max_patch = major, minor, patch
    except Exception as e:
        logger.warning(f"自动识别文件名版本号失败: {e}")

    version = "1.0.0"
    if found:
        version = f"{max_major}.{max_minor}.{max_patch + 1}"

    return f"{safe_uni}{requirements}_V{version}.{ext}"


def _build_rows(data: list[dict]) -> list[list]:
    rows = []
    for i, row in enumerate(data, 1):
        rows.append([
            i,
            row.get("name") or "",
            row.get("email") or "",
            row.get("department") or "",
            row.get("title") or "",
            row.get("url") or "",
        ])
    return rows


# —————————————————————————— CSV ——————————————————————————


def export_csv(data: list[dict], university: str, task_id: str = "") -> Path:
    out_dir = get_task_dir(task_id)
    filepath = out_dir / _make_filename(university, "csv", task_id)

    with open(filepath, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(HEADERS)
        for row in _build_rows(data):
            writer.writerow(row)

    logger.info(f"CSV 已保存: {filepath}")
    return filepath


# —————————————————————————— XLSX ——————————————————————————


def export_xlsx(data: list[dict], university: str, task_id: str = "") -> Path:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    out_dir = get_task_dir(task_id)
    filepath = out_dir / _make_filename(university, "xlsx", task_id)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "教师邮箱"

    header_font = Font(name="微软雅黑", bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="10A37F", end_color="10A37F", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center")
    thin_border = Border(
        left=Side(style="thin", color="D1D5DB"),
        right=Side(style="thin", color="D1D5DB"),
        top=Side(style="thin", color="D1D5DB"),
        bottom=Side(style="thin", color="D1D5DB"),
    )
    cell_font = Font(name="微软雅黑", size=10)
    cell_align = Alignment(vertical="center")

    for col, header in enumerate(HEADERS, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    for i, row in enumerate(_build_rows(data), 1):
        for col, val in enumerate(row, 1):
            cell = ws.cell(row=i + 1, column=col, value=val)
            cell.font = cell_font
            cell.alignment = cell_align
            cell.border = thin_border

    ws.column_dimensions["A"].width = 8
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 32
    ws.column_dimensions["D"].width = 20
    ws.column_dimensions["E"].width = 16
    ws.column_dimensions["F"].width = 55

    wb.save(filepath)
    logger.info(f"XLSX 已保存: {filepath}")
    return filepath


# —————————————————————————— Markdown ——————————————————————————


def export_markdown(data: list[dict], university: str, task_id: str = "") -> Path:
    out_dir = get_task_dir(task_id)
    filepath = out_dir / _make_filename(university, "md", task_id)

    lines = [
        f"# {university} — 教师邮箱列表",
        "",
        f"**抓取时间**：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"**记录数量**：{len(data)} 条",
        "",
        "| 序号 | 姓名 | 邮箱 | 学院 | 职称 | 主页链接 |",
        "|------|------|------|------|------|----------|",
    ]

    for i, row in enumerate(data, 1):
        name = str(row.get("name", "")).replace("|", "\\|")
        email = str(row.get("email", "")).replace("|", "\\|")
        dept = str(row.get("department", "")).replace("|", "\\|")
        title = str(row.get("title", "")).replace("|", "\\|")
        url = str(row.get("url", "")).replace("|", "\\|")
        lines.append(f"| {i} | {name} | {email} | {dept} | {title} | {url} |")

    lines.append("")
    lines.append(f"> 共 {len(data)} 条记录，由 UniEmail Agent 自动生成。")

    with open(filepath, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    logger.info(f"Markdown 已保存: {filepath}")
    return filepath


# —————————————————————————— HTML ——————————————————————————


def export_html(data: list[dict], university: str, task_id: str = "") -> Path:
    out_dir = get_task_dir(task_id)
    filepath = out_dir / _make_filename(university, "html", task_id)

    rows_html = ""
    for i, row in enumerate(data, 1):
        rows_html += f"""        <tr>
            <td>{i}</td>
            <td>{_escape_html(row.get("name", ""))}</td>
            <td><a href="mailto:{_escape_html(row.get("email", ""))}">{_escape_html(row.get("email", ""))}</a></td>
            <td>{_escape_html(row.get("department", ""))}</td>
            <td>{_escape_html(row.get("title", ""))}</td>
            <td><a href="{_escape_html(row.get("url", ""))}" target="_blank">{_escape_html(row.get("url", ""))}</a></td>
        </tr>\n"""

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{_escape_html(university)} — 教师邮箱列表</title>
<style>
  :root {{ --primary: #10A37F; --border: #E5E7EB; }}
  body {{ font-family: -apple-system, "Microsoft YaHei", sans-serif; max-width: 1200px; margin: 0 auto; padding: 40px 24px; color: #212121; background: #FAFAFA; }}
  h1 {{ color: var(--primary); font-size: 1.5rem; margin-bottom: 8px; }}
  .meta {{ color: #6E6E80; font-size: 0.875rem; margin-bottom: 24px; }}
  table {{ width: 100%; border-collapse: collapse; background: #FFF; border-radius: 12px; overflow: hidden; box-shadow: 0 1px 3px rgba(0,0,0,0.06); }}
  th {{ background: var(--primary); color: #FFF; padding: 12px 16px; text-align: left; font-size: 0.875rem; font-weight: 600; }}
  td {{ padding: 10px 16px; border-bottom: 1px solid var(--border); font-size: 0.875rem; }}
  tr:last-child td {{ border-bottom: none; }}
  tr:hover td {{ background: #F5F7F8; }}
  a {{ color: var(--primary); text-decoration: none; }}
  a:hover {{ text-decoration: underline; }}
  footer {{ margin-top: 24px; color: #9A9AA5; font-size: 0.75rem; }}
</style>
</head>
<body>
<h1>{_escape_html(university)} — 教师邮箱列表</h1>
<p class="meta">生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} &nbsp;|&nbsp; 记录数量：{len(data)} 条</p>
<table>
    <thead>
        <tr><th>序号</th><th>姓名</th><th>邮箱</th><th>学院</th><th>职称</th><th>主页链接</th></tr>
    </thead>
    <tbody>
{rows_html}    </tbody>
</table>
<footer>由 UniEmail Agent 自动生成</footer>
</body>
</html>"""

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html)

    logger.info(f"HTML 已保存: {filepath}")
    return filepath


def _escape_html(text: str) -> str:
    return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


# —————————————————————————— PDF ——————————————————————————


def _find_cjk_font() -> str | None:
    import platform
    candidates = []
    if platform.system() == "Windows":
        candidates = [
            "C:/Windows/Fonts/msyh.ttc",
            "C:/Windows/Fonts/simsun.ttc",
            "C:/Windows/Fonts/msyhbd.ttf",
        ]
    else:
        candidates = [
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
        ]
    for path in candidates:
        if Path(path).exists():
            return path
    return None


def export_pdf(data: list[dict], university: str, task_id: str = "") -> Path:
    from fpdf import FPDF

    out_dir = get_task_dir(task_id)
    filepath = out_dir / _make_filename(university, "pdf", task_id)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    font_path = _find_cjk_font()
    use_cjk = font_path is not None
    if use_cjk:
        pdf.add_font("cjk", "", font_path, uni=True)
        pdf.add_font("cjk", "B", font_path, uni=True)
    else:
        logger.warning("未找到 CJK 字体，PDF 中文可能显示异常")

    pdf.add_page()

    if use_cjk:
        pdf.set_font("cjk", "B", 16)
    else:
        pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(16, 163, 127)
    pdf.cell(0, 10, f"{university} — 教师邮箱列表", align="L")
    pdf.ln(8)

    if use_cjk:
        pdf.set_font("cjk", "", 9)
    else:
        pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(110, 110, 128)
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    pdf.cell(0, 6, f"生成时间: {ts}    记录数量: {len(data)} 条", align="L")
    pdf.ln(8)

    col_widths = [10, 32, 55, 32, 24, 42]
    col_labels = ["#", "姓名", "邮箱", "学院", "职称", "主页链接"]

    pdf.set_fill_color(16, 163, 127)
    pdf.set_text_color(255, 255, 255)
    if use_cjk:
        pdf.set_font("cjk", "B", 8)
    else:
        pdf.set_font("Helvetica", "B", 8)

    for w, label in zip(col_widths, col_labels):
        pdf.cell(w, 8, label, border=1, fill=True, align="C")
    pdf.ln()

    if use_cjk:
        pdf.set_font("cjk", "", 8)
    else:
        pdf.set_font("Helvetica", "", 7)
    pdf.set_text_color(33, 33, 33)

    row_fill = False
    for i, row in enumerate(data, 1):
        if pdf.get_y() > 265:
            pdf.add_page()
            pdf.set_fill_color(16, 163, 127)
            pdf.set_text_color(255, 255, 255)
            if use_cjk:
                pdf.set_font("cjk", "B", 8)
            else:
                pdf.set_font("Helvetica", "B", 8)
            for w, label in zip(col_widths, col_labels):
                pdf.cell(w, 8, label, border=1, fill=True, align="C")
            pdf.ln()
            if use_cjk:
                pdf.set_font("cjk", "", 8)
            else:
                pdf.set_font("Helvetica", "", 7)
            pdf.set_text_color(33, 33, 33)

        if row_fill:
            pdf.set_fill_color(245, 247, 248)
        else:
            pdf.set_fill_color(255, 255, 255)
        row_fill = not row_fill

        values = [
            str(i),
            str(row.get("name", ""))[:20],
            str(row.get("email", ""))[:40],
            str(row.get("department", ""))[:18],
            str(row.get("title", ""))[:14],
            str(row.get("url", ""))[:30],
        ]
        for w, val in zip(col_widths, values):
            pdf.cell(w, 7, val, border=1, fill=True)
        pdf.ln()

    pdf.output(str(filepath))
    logger.info(f"PDF 已保存: {filepath}")
    return filepath


# —————————————————————————— DOCX ——————————————————————————


def export_docx(data: list[dict], university: str, task_id: str = "") -> Path:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    out_dir = get_task_dir(task_id)
    filepath = out_dir / _make_filename(university, "docx", task_id)

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    title = doc.add_heading(f"{university} — 教师邮箱列表", level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x10, 0xA3, 0x7F)

    meta = doc.add_paragraph()
    meta.style = doc.styles["Normal"]
    meta_run = meta.add_run(
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}    记录数量：{len(data)} 条"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x6E, 0x6E, 0x80)

    table = doc.add_table(rows=1 + len(data), cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for i, label in enumerate(HEADERS):
        cell = header_cells[i]
        cell.text = str(label)
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "10A37F",
            qn("w:val"): "clear",
        })
        shading.append(shading_elm)
        for p in cell.paragraphs:
            p.alignment = 1
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)

    for i, row in enumerate(data, 1):
        row_cells = table.rows[i].cells
        values = [
            str(i),
            str(row.get("name", "")),
            str(row.get("email", "")),
            str(row.get("department", "")),
            str(row.get("title", "")),
            str(row.get("url", "")),
        ]
        for j, val in enumerate(values):
            row_cells[j].text = val
            for p in row_cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)

    footer = doc.add_paragraph()
    footer.alignment = 2
    footer_run = footer.add_run("由 UniEmail Agent 自动生成")
    footer_run.font.size = Pt(7.5)
    footer_run.font.color.rgb = RGBColor(0x9A, 0x9A, 0xA5)

    doc.save(str(filepath))
    logger.info(f"DOCX 已保存: {filepath}")
    return filepath


# —————————————————————————— 批量导出 ——————————————————————————


def export_all(
    data: list[dict], university: str, task_id: str = "",
    formats: list[str] | None = None,
) -> dict:
    """导出指定格式，返回 {format: filename}。formats=None 时导出全部。"""
    exporters = {
        "csv": export_csv,
        "xlsx": export_xlsx,
        "md": export_markdown,
        "html": export_html,
        "pdf": export_pdf,
        "docx": export_docx,
    }
    keys = formats or list(exporters)
    result = {}
    for fmt in keys:
        fn = exporters.get(fmt)
        if fn:
            try:
                result[fmt] = fn(data, university, task_id).name
            except Exception as e:
                logger.error(f"{fmt.upper()} 导出失败: {e}")
    return result